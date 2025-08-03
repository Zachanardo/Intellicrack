"""
Report Format Generators for AI Report Generation System

This module provides specialized generators for different report output formats
including HTML, PDF, Markdown, JSON, DOCX, and plain text.

Copyright (C) 2025 Zachary Flint

This file is part of Intellicrack.

Intellicrack is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

Intellicrack is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with Intellicrack.  If not, see <https://www.gnu.org/licenses/>.
"""

import base64
import json
import logging
import os
import tempfile
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from .ai_report_generator import ReportFinding, ReportMetadata, AIInsight, ReportFormat, ReportSeverity
from ...utils.logger import get_logger

logger = get_logger(__name__)

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.piecharts import Pie
    PDF_AVAILABLE = True
except ImportError:
    logger.warning("ReportLab not available - PDF generation will use fallback")
    PDF_AVAILABLE = False

try:
    import docx
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available - DOCX generation will be skipped")
    DOCX_AVAILABLE = False

try:
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    logger.warning("Matplotlib not available - charts will be skipped")
    MATPLOTLIB_AVAILABLE = False


class BaseFormatGenerator(ABC):
    """Base class for report format generators."""
    
    def __init__(self, output_dir: str = "reports"):
        """Initialize the format generator."""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    @abstractmethod
    async def generate(self, content: Dict[str, Any], output_filename: str) -> str:
        """Generate report in specific format."""
        pass
    
    def _get_severity_color(self, severity: ReportSeverity) -> str:
        """Get color code for severity level."""
        colors = {
            ReportSeverity.CRITICAL: "#8B0000",  # Dark red
            ReportSeverity.HIGH: "#FF4500",      # Orange red
            ReportSeverity.MEDIUM: "#FFD700",    # Gold
            ReportSeverity.LOW: "#32CD32",       # Lime green
            ReportSeverity.INFO: "#87CEEB"       # Sky blue
        }
        return colors.get(severity, "#808080")  # Gray default
    
    def _get_severity_priority(self, severity: ReportSeverity) -> int:
        """Get numeric priority for severity (lower = higher priority)."""
        priorities = {
            ReportSeverity.CRITICAL: 1,
            ReportSeverity.HIGH: 2,
            ReportSeverity.MEDIUM: 3,
            ReportSeverity.LOW: 4,
            ReportSeverity.INFO: 5
        }
        return priorities.get(severity, 6)
    
    def _format_timestamp(self, timestamp: datetime) -> str:
        """Format timestamp for display."""
        return timestamp.strftime("%Y-%m-%d %H:%M:%S UTC")


class HTMLGenerator(BaseFormatGenerator):
    """Generate HTML reports with interactive features."""
    
    async def generate(self, content: Dict[str, Any], output_filename: str) -> str:
        """Generate HTML report."""
        output_path = self.output_dir / f"{output_filename}.html"
        
        html_content = await self._create_html_content(content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"HTML report generated: {output_path}")
        return str(output_path)
    
    async def _create_html_content(self, content: Dict[str, Any]) -> str:
        """Create complete HTML content."""
        metadata = content.get('metadata')
        findings = content.get('findings', [])
        ai_insights = content.get('ai_insights', [])
        
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.title if metadata else 'Security Analysis Report'}</title>
    <style>
        {self._get_html_styles()}
    </style>
    <script>
        {self._get_html_javascript()}
    </script>
</head>
<body>
    <div class="container">
        {await self._create_html_header(metadata)}
        {await self._create_html_navigation()}
        
        <main>
            {await self._create_html_executive_summary(content)}
            {await self._create_html_findings_overview(findings)}
            {await self._create_html_ai_insights(ai_insights)}
            {await self._create_html_detailed_findings(findings)}
            {await self._create_html_risk_assessment(content)}
            {await self._create_html_recommendations(content)}
            {await self._create_html_appendices(content)}
        </main>
        
        {await self._create_html_footer()}
    </div>
</body>
</html>
"""
        return html
    
    def _get_html_styles(self) -> str:
        """Get CSS styles for HTML report."""
        return """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #f5f5f5;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background-color: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }
        
        header {
            background: linear-gradient(135deg, #2c3e50, #3498db);
            color: white;
            padding: 2rem;
            text-align: center;
        }
        
        header h1 {
            font-size: 2.5rem;
            margin-bottom: 0.5rem;
        }
        
        header .subtitle {
            font-size: 1.2rem;
            opacity: 0.9;
        }
        
        nav {
            background-color: #34495e;
            padding: 1rem;
            position: sticky;
            top: 0;
            z-index: 100;
        }
        
        nav ul {
            list-style: none;
            display: flex;
            justify-content: center;
            flex-wrap: wrap;
        }
        
        nav li {
            margin: 0 1rem;
        }
        
        nav a {
            color: white;
            text-decoration: none;
            padding: 0.5rem 1rem;
            border-radius: 4px;
            transition: background-color 0.3s;
        }
        
        nav a:hover {
            background-color: rgba(255,255,255,0.2);
        }
        
        main {
            padding: 2rem;
        }
        
        .section {
            margin-bottom: 3rem;
            padding: 1.5rem;
            border-radius: 8px;
            background-color: #fff;
            border-left: 4px solid #3498db;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }
        
        .section h2 {
            color: #2c3e50;
            margin-bottom: 1rem;
            font-size: 1.8rem;
            border-bottom: 2px solid #ecf0f1;
            padding-bottom: 0.5rem;
        }
        
        .section h3 {
            color: #34495e;
            margin: 1.5rem 0 1rem 0;
            font-size: 1.4rem;
        }
        
        .findings-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 1rem;
            margin: 1rem 0;
        }
        
        .finding-card {
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 1rem;
            background-color: #fafafa;
            border-left: 4px solid #ccc;
        }
        
        .finding-card.critical {
            border-left-color: #8B0000;
            background-color: #fff5f5;
        }
        
        .finding-card.high {
            border-left-color: #FF4500;
            background-color: #fff8f0;
        }
        
        .finding-card.medium {
            border-left-color: #FFD700;
            background-color: #fffef0;
        }
        
        .finding-card.low {
            border-left-color: #32CD32;
            background-color: #f5fff5;
        }
        
        .finding-card.info {
            border-left-color: #87CEEB;
            background-color: #f0f8ff;
        }
        
        .severity-badge {
            display: inline-block;
            padding: 0.25rem 0.75rem;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: bold;
            text-transform: uppercase;
            color: white;
        }
        
        .severity-critical { background-color: #8B0000; }
        .severity-high { background-color: #FF4500; }
        .severity-medium { background-color: #FFD700; color: #333; }
        .severity-low { background-color: #32CD32; }
        .severity-info { background-color: #87CEEB; }
        
        .ai-insight {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 1.5rem;
            border-radius: 8px;
            margin: 1rem 0;
        }
        
        .ai-insight h4 {
            margin-bottom: 0.5rem;
            font-size: 1.2rem;
        }
        
        .confidence-bar {
            width: 100%;
            height: 8px;
            background-color: rgba(255,255,255,0.3);
            border-radius: 4px;
            overflow: hidden;
            margin-top: 0.5rem;
        }
        
        .confidence-fill {
            height: 100%;
            background-color: rgba(255,255,255,0.9);
            transition: width 0.3s ease;
        }
        
        .stats-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1rem;
            margin: 1rem 0;
        }
        
        .stat-card {
            text-align: center;
            padding: 1.5rem;
            background: linear-gradient(135deg, #74b9ff, #0984e3);
            color: white;
            border-radius: 8px;
        }
        
        .stat-value {
            font-size: 2.5rem;
            font-weight: bold;
            display: block;
        }
        
        .stat-label {
            font-size: 0.9rem;
            opacity: 0.9;
        }
        
        .table-responsive {
            overflow-x: auto;
            margin: 1rem 0;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 1rem 0;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 0.75rem;
            text-align: left;
        }
        
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        
        .chart-container {
            text-align: center;
            margin: 2rem 0;
        }
        
        footer {
            background-color: #2c3e50;
            color: white;
            text-align: center;
            padding: 2rem;
        }
        
        .timestamp {
            color: #666;
            font-size: 0.9rem;
            font-style: italic;
        }
        
        .expandable {
            cursor: pointer;
            user-select: none;
        }
        
        .expandable:hover {
            background-color: #f0f0f0;
        }
        
        .collapsible-content {
            max-height: 0;
            overflow: hidden;
            transition: max-height 0.3s ease;
        }
        
        .collapsible-content.expanded {
            max-height: 1000px;
        }
        
        @media (max-width: 768px) {
            .container {
                margin: 0;
            }
            
            header {
                padding: 1rem;
            }
            
            header h1 {
                font-size: 2rem;
            }
            
            main {
                padding: 1rem;
            }
            
            nav ul {
                flex-direction: column;
                align-items: center;
            }
            
            nav li {
                margin: 0.25rem 0;
            }
        }
        """
    
    def _get_html_javascript(self) -> str:
        """Get JavaScript for HTML report interactivity."""
        return """
        function toggleCollapsible(element) {
            const content = element.nextElementSibling;
            content.classList.toggle('expanded');
        }
        
        function showSection(sectionId) {
            const section = document.getElementById(sectionId);
            if (section) {
                section.scrollIntoView({ behavior: 'smooth' });
            }
        }
        
        function filterFindings(severity) {
            const cards = document.querySelectorAll('.finding-card');
            cards.forEach(card => {
                if (severity === 'all' || card.classList.contains(severity)) {
                    card.style.display = 'block';
                } else {
                    card.style.display = 'none';
                }
            });
        }
        
        document.addEventListener('DOMContentLoaded', function() {
            // Initialize confidence bars animation
            const confidenceBars = document.querySelectorAll('.confidence-fill');
            confidenceBars.forEach(bar => {
                const width = bar.getAttribute('data-confidence');
                setTimeout(() => {
                    bar.style.width = width + '%';
                }, 500);
            });
        });
        """
    
    async def _create_html_header(self, metadata: Optional[ReportMetadata]) -> str:
        """Create HTML header section."""
        if not metadata:
            return "<header><h1>Security Analysis Report</h1></header>"
        
        return f"""
        <header>
            <h1>{metadata.title}</h1>
            <div class="subtitle">
                Generated on {self._format_timestamp(metadata.generated_at)}<br>
                Target: {Path(metadata.target_binary).name}
            </div>
        </header>
        """
    
    async def _create_html_navigation(self) -> str:
        """Create HTML navigation menu."""
        return """
        <nav>
            <ul>
                <li><a href="#executive-summary" onclick="showSection('executive-summary')">Executive Summary</a></li>
                <li><a href="#findings-overview" onclick="showSection('findings-overview')">Findings Overview</a></li>
                <li><a href="#ai-insights" onclick="showSection('ai-insights')">AI Insights</a></li>
                <li><a href="#detailed-findings" onclick="showSection('detailed-findings')">Detailed Findings</a></li>
                <li><a href="#risk-assessment" onclick="showSection('risk-assessment')">Risk Assessment</a></li>
                <li><a href="#recommendations" onclick="showSection('recommendations')">Recommendations</a></li>
            </ul>
        </nav>
        """
    
    async def _create_html_executive_summary(self, content: Dict[str, Any]) -> str:
        """Create HTML executive summary section."""
        summary = content.get('executive_summary', 'No executive summary available.')
        
        return f"""
        <section id="executive-summary" class="section">
            <h2>Executive Summary</h2>
            <p>{summary}</p>
        </section>
        """
    
    async def _create_html_findings_overview(self, findings: List[ReportFinding]) -> str:
        """Create HTML findings overview section."""
        if not findings:
            return """
            <section id="findings-overview" class="section">
                <h2>Findings Overview</h2>
                <p>No findings to report.</p>
            </section>
            """
        
        # Calculate statistics
        severity_counts = {}
        for severity in ReportSeverity:
            severity_counts[severity] = len([f for f in findings if f.severity == severity])
        
        stats_html = ""
        for severity, count in severity_counts.items():
            if count > 0:
                stats_html += f"""
                <div class="stat-card">
                    <span class="stat-value">{count}</span>
                    <span class="stat-label">{severity.value.title()}</span>
                </div>
                """
        
        # Create severity filter
        filter_html = """
        <div style="margin: 1rem 0;">
            <strong>Filter by severity:</strong>
            <button onclick="filterFindings('all')" style="margin: 0 0.25rem; padding: 0.25rem 0.5rem;">All</button>
            <button onclick="filterFindings('critical')" style="margin: 0 0.25rem; padding: 0.25rem 0.5rem;">Critical</button>
            <button onclick="filterFindings('high')" style="margin: 0 0.25rem; padding: 0.25rem 0.5rem;">High</button>
            <button onclick="filterFindings('medium')" style="margin: 0 0.25rem; padding: 0.25rem 0.5rem;">Medium</button>
            <button onclick="filterFindings('low')" style="margin: 0 0.25rem; padding: 0.25rem 0.5rem;">Low</button>
            <button onclick="filterFindings('info')" style="margin: 0 0.25rem; padding: 0.25rem 0.5rem;">Info</button>
        </div>
        """
        
        return f"""
        <section id="findings-overview" class="section">
            <h2>Findings Overview</h2>
            <div class="stats-grid">
                {stats_html}
            </div>
            {filter_html}
            <p>Total findings: <strong>{len(findings)}</strong></p>
        </section>
        """
    
    async def _create_html_ai_insights(self, ai_insights: List[AIInsight]) -> str:
        """Create HTML AI insights section."""
        if not ai_insights:
            return """
            <section id="ai-insights" class="section">
                <h2>AI Insights</h2>
                <p>No AI insights available.</p>
            </section>
            """
        
        insights_html = ""
        for insight in ai_insights:
            insights_html += f"""
            <div class="ai-insight">
                <h4>{insight.title}</h4>
                <p>{insight.content}</p>
                <div class="confidence-bar">
                    <div class="confidence-fill" data-confidence="{insight.confidence * 100}"></div>
                </div>
                <small>Confidence: {insight.confidence:.1%} | Model: {insight.model_used}</small>
            </div>
            """
        
        return f"""
        <section id="ai-insights" class="section">
            <h2>AI-Generated Insights</h2>
            {insights_html}
        </section>
        """
    
    async def _create_html_detailed_findings(self, findings: List[ReportFinding]) -> str:
        """Create HTML detailed findings section."""
        if not findings:
            return """
            <section id="detailed-findings" class="section">
                <h2>Detailed Findings</h2>
                <p>No detailed findings to report.</p>
            </section>
            """
        
        # Sort findings by severity
        sorted_findings = sorted(findings, key=lambda f: self._get_severity_priority(f.severity))
        
        findings_html = ""
        for finding in sorted_findings:
            evidence_html = ""
            if finding.evidence:
                evidence_html = "<ul>" + "".join(f"<li>{evidence}</li>" for evidence in finding.evidence) + "</ul>"
            
            recommendations_html = ""
            if finding.recommendations:
                recommendations_html = "<ul>" + "".join(f"<li>{rec}</li>" for rec in finding.recommendations) + "</ul>"
            
            findings_html += f"""
            <div class="finding-card {finding.severity.value}">
                <h4>{finding.title} <span class="severity-badge severity-{finding.severity.value}">{finding.severity.value}</span></h4>
                <p><strong>Category:</strong> {finding.category}</p>
                <p><strong>Description:</strong> {finding.description}</p>
                
                {f"<p><strong>Evidence:</strong></p>{evidence_html}" if evidence_html else ""}
                {f"<p><strong>Recommendations:</strong></p>{recommendations_html}" if recommendations_html else ""}
                
                <div style="margin-top: 1rem; font-size: 0.9rem; color: #666;">
                    <strong>Confidence:</strong> {finding.confidence:.1%} |
                    <strong>Exploitation Difficulty:</strong> {finding.exploitation_difficulty} |
                    <strong>Business Impact:</strong> {finding.business_impact}
                </div>
            </div>
            """
        
        return f"""
        <section id="detailed-findings" class="section">
            <h2>Detailed Findings</h2>
            <div class="findings-grid">
                {findings_html}
            </div>
        </section>
        """
    
    async def _create_html_risk_assessment(self, content: Dict[str, Any]) -> str:
        """Create HTML risk assessment section."""
        risk_assessment = content.get('risk_assessment', 'No risk assessment available.')
        
        return f"""
        <section id="risk-assessment" class="section">
            <h2>Risk Assessment</h2>
            <div>{risk_assessment}</div>
        </section>
        """
    
    async def _create_html_recommendations(self, content: Dict[str, Any]) -> str:
        """Create HTML recommendations section."""
        recommendations = content.get('recommendations', 'No recommendations available.')
        
        return f"""
        <section id="recommendations" class="section">
            <h2>Recommendations</h2>
            <div>{recommendations}</div>
        </section>
        """
    
    async def _create_html_appendices(self, content: Dict[str, Any]) -> str:
        """Create HTML appendices section."""
        metadata = content.get('metadata')
        if not metadata:
            return ""
        
        return f"""
        <section id="appendices" class="section">
            <h2>Appendices</h2>
            <h3>Report Metadata</h3>
            <table>
                <tr><td><strong>Report ID</strong></td><td>{metadata.report_id}</td></tr>
                <tr><td><strong>Generated By</strong></td><td>{metadata.generated_by}</td></tr>
                <tr><td><strong>Version</strong></td><td>{metadata.version}</td></tr>
                <tr><td><strong>Analysis Duration</strong></td><td>{metadata.analysis_duration or 'Unknown'}</td></tr>
                <tr><td><strong>Data Sources</strong></td><td>{', '.join(metadata.data_sources) if metadata.data_sources else 'None'}</td></tr>
            </table>
        </section>
        """
    
    async def _create_html_footer(self) -> str:
        """Create HTML footer."""
        return f"""
        <footer>
            <p>Generated by Intellicrack AI Report Generator</p>
            <p class="timestamp">Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}</p>
        </footer>
        """


class PDFGenerator(BaseFormatGenerator):
    """Generate PDF reports using ReportLab."""
    
    async def generate(self, content: Dict[str, Any], output_filename: str) -> str:
        """Generate PDF report."""
        if not PDF_AVAILABLE:
            # Fallback to HTML
            html_gen = HTMLGenerator(str(self.output_dir))
            html_path = await html_gen.generate(content, output_filename)
            logger.warning(f"PDF generation not available, generated HTML instead: {html_path}")
            return html_path
        
        output_path = self.output_dir / f"{output_filename}.pdf"
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center
            textColor=colors.HexColor('#2c3e50')
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#34495e')
        )
        
        # Add content
        await self._add_pdf_header(story, content, title_style)
        await self._add_pdf_executive_summary(story, content, styles, heading_style)
        await self._add_pdf_findings_overview(story, content, styles, heading_style)
        await self._add_pdf_detailed_findings(story, content, styles, heading_style)
        await self._add_pdf_ai_insights(story, content, styles, heading_style)
        await self._add_pdf_recommendations(story, content, styles, heading_style)
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"PDF report generated: {output_path}")
        return str(output_path)
    
    async def _add_pdf_header(self, story: List, content: Dict[str, Any], title_style):
        """Add PDF header."""
        metadata = content.get('metadata')
        
        if metadata:
            story.append(Paragraph(metadata.title, title_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"Generated: {self._format_timestamp(metadata.generated_at)}", styles['Normal']))
            story.append(Paragraph(f"Target: {Path(metadata.target_binary).name}", styles['Normal']))
        else:
            story.append(Paragraph("Security Analysis Report", title_style))
        
        story.append(Spacer(1, 24))
    
    async def _add_pdf_executive_summary(self, story: List, content: Dict[str, Any], styles, heading_style):
        """Add executive summary to PDF."""
        story.append(Paragraph("Executive Summary", heading_style))
        summary = content.get('executive_summary', 'No executive summary available.')
        story.append(Paragraph(summary, styles['Normal']))
        story.append(Spacer(1, 12))
    
    async def _add_pdf_findings_overview(self, story: List, content: Dict[str, Any], styles, heading_style):
        """Add findings overview to PDF."""
        findings = content.get('findings', [])
        
        story.append(Paragraph("Findings Overview", heading_style))
        
        if not findings:
            story.append(Paragraph("No findings to report.", styles['Normal']))
            story.append(Spacer(1, 12))
            return
        
        # Create severity summary table
        severity_counts = {}
        for severity in ReportSeverity:
            severity_counts[severity] = len([f for f in findings if f.severity == severity])
        
        table_data = [['Severity', 'Count']]
        for severity, count in severity_counts.items():
            if count > 0:
                table_data.append([severity.value.title(), str(count)])
        
        if len(table_data) > 1:
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
        
        story.append(Spacer(1, 12))
    
    async def _add_pdf_detailed_findings(self, story: List, content: Dict[str, Any], styles, heading_style):
        """Add detailed findings to PDF."""
        findings = content.get('findings', [])
        
        story.append(Paragraph("Detailed Findings", heading_style))
        
        if not findings:
            story.append(Paragraph("No detailed findings to report.", styles['Normal']))
            story.append(Spacer(1, 12))
            return
        
        # Sort findings by severity
        sorted_findings = sorted(findings, key=lambda f: self._get_severity_priority(f.severity))
        
        for finding in sorted_findings:
            # Finding header
            story.append(Paragraph(f"{finding.title} ({finding.severity.value.upper()})", styles['Heading3']))
            
            # Finding details
            story.append(Paragraph(f"<b>Category:</b> {finding.category}", styles['Normal']))
            story.append(Paragraph(f"<b>Description:</b> {finding.description}", styles['Normal']))
            
            if finding.evidence:
                story.append(Paragraph("<b>Evidence:</b>", styles['Normal']))
                for evidence in finding.evidence:
                    story.append(Paragraph(f"• {evidence}", styles['Normal']))
            
            if finding.recommendations:
                story.append(Paragraph("<b>Recommendations:</b>", styles['Normal']))
                for rec in finding.recommendations:
                    story.append(Paragraph(f"• {rec}", styles['Normal']))
            
            story.append(Paragraph(
                f"<b>Confidence:</b> {finding.confidence:.1%} | "
                f"<b>Exploitation Difficulty:</b> {finding.exploitation_difficulty} | "
                f"<b>Business Impact:</b> {finding.business_impact}",
                styles['Normal']
            ))
            
            story.append(Spacer(1, 12))
    
    async def _add_pdf_ai_insights(self, story: List, content: Dict[str, Any], styles, heading_style):
        """Add AI insights to PDF."""
        ai_insights = content.get('ai_insights', [])
        
        story.append(Paragraph("AI-Generated Insights", heading_style))
        
        if not ai_insights:
            story.append(Paragraph("No AI insights available.", styles['Normal']))
            story.append(Spacer(1, 12))
            return
        
        for insight in ai_insights:
            story.append(Paragraph(insight.title, styles['Heading4']))
            story.append(Paragraph(insight.content, styles['Normal']))
            story.append(Paragraph(
                f"<i>Confidence: {insight.confidence:.1%} | Model: {insight.model_used}</i>",
                styles['Normal']
            ))
            story.append(Spacer(1, 8))
    
    async def _add_pdf_recommendations(self, story: List, content: Dict[str, Any], styles, heading_style):
        """Add recommendations to PDF."""
        story.append(Paragraph("Recommendations", heading_style))
        recommendations = content.get('recommendations', 'No recommendations available.')
        story.append(Paragraph(recommendations, styles['Normal']))


class MarkdownGenerator(BaseFormatGenerator):
    """Generate Markdown reports."""
    
    async def generate(self, content: Dict[str, Any], output_filename: str) -> str:
        """Generate Markdown report."""
        output_path = self.output_dir / f"{output_filename}.md"
        
        markdown_content = await self._create_markdown_content(content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        logger.info(f"Markdown report generated: {output_path}")
        return str(output_path)    
    async def _create_markdown_content(self, content: Dict[str, Any]) -> str:
        """Create complete Markdown content."""
        metadata = content.get('metadata')
        findings = content.get('findings', [])
        ai_insights = content.get('ai_insights', [])
        
        lines = []
        
        # Header
        if metadata:
            lines.extend([
                f"# {metadata.title}",
                "",
                f"**Generated:** {self._format_timestamp(metadata.generated_at)}  ",
                f"**Target:** {Path(metadata.target_binary).name}  ",
                f"**Report ID:** {metadata.report_id}  ",
                ""
            ])
        else:
            lines.extend(["# Security Analysis Report", ""])
        
        # Table of Contents
        lines.extend([
            "## Table of Contents",
            "",
            "- [Executive Summary](#executive-summary)",
            "- [Findings Overview](#findings-overview)",
            "- [AI Insights](#ai-insights)",
            "- [Detailed Findings](#detailed-findings)",
            "- [Risk Assessment](#risk-assessment)",
            "- [Recommendations](#recommendations)",
            "- [Appendices](#appendices)",
            ""
        ])
        
        # Executive Summary
        lines.extend([
            "## Executive Summary",
            "",
            content.get('executive_summary', 'No executive summary available.'),
            ""
        ])
        
        # Findings Overview
        lines.append("## Findings Overview")
        lines.append("")
        
        if findings:
            severity_counts = {}
            for severity in ReportSeverity:
                severity_counts[severity] = len([f for f in findings if f.severity == severity])
            
            lines.append("| Severity | Count |")
            lines.append("|----------|-------|")
            
            for severity, count in severity_counts.items():
                if count > 0:
                    lines.append(f"| {severity.value.title()} | {count} |")
            
            lines.extend(["", f"**Total Findings:** {len(findings)}", ""])
        else:
            lines.extend(["No findings to report.", ""])
        
        # AI Insights
        lines.append("## AI Insights")
        lines.append("")
        
        if ai_insights:
            for insight in ai_insights:
                lines.extend([
                    f"### {insight.title}",
                    "",
                    insight.content,
                    "",
                    f"*Confidence: {insight.confidence:.1%} | Model: {insight.model_used}*",
                    ""
                ])
        else:
            lines.extend(["No AI insights available.", ""])
        
        # Detailed Findings
        lines.append("## Detailed Findings")
        lines.append("")
        
        if findings:
            sorted_findings = sorted(findings, key=lambda f: self._get_severity_priority(f.severity))
            
            for finding in sorted_findings:
                lines.extend([
                    f"### {finding.title}",
                    "",
                    f"**Severity:** {finding.severity.value.upper()}  ",
                    f"**Category:** {finding.category}  ",
                    f"**Description:** {finding.description}  ",
                    ""
                ])
                
                if finding.evidence:
                    lines.append("**Evidence:**")
                    for evidence in finding.evidence:
                        lines.append(f"- {evidence}")
                    lines.append("")
                
                if finding.recommendations:
                    lines.append("**Recommendations:**")
                    for rec in finding.recommendations:
                        lines.append(f"- {rec}")
                    lines.append("")
                
                lines.extend([
                    f"**Confidence:** {finding.confidence:.1%} | "
                    f"**Exploitation Difficulty:** {finding.exploitation_difficulty} | "
                    f"**Business Impact:** {finding.business_impact}",
                    "",
                    "---",
                    ""
                ])
        else:
            lines.extend(["No detailed findings to report.", ""])
        
        # Risk Assessment
        lines.extend([
            "## Risk Assessment",
            "",
            content.get('risk_assessment', 'No risk assessment available.'),
            ""
        ])
        
        # Recommendations
        lines.extend([
            "## Recommendations",
            "",
            content.get('recommendations', 'No recommendations available.'),
            ""
        ])
        
        # Appendices
        lines.append("## Appendices")
        lines.append("")
        
        if metadata:
            lines.extend([
                "### Report Metadata",
                "",
                f"- **Report ID:** {metadata.report_id}",
                f"- **Generated By:** {metadata.generated_by}",
                f"- **Version:** {metadata.version}",
                f"- **Analysis Duration:** {metadata.analysis_duration or 'Unknown'}",
                f"- **Data Sources:** {', '.join(metadata.data_sources) if metadata.data_sources else 'None'}",
                ""
            ])
        
        lines.extend([
            "---",
            f"*Report generated by Intellicrack AI Report Generator on {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}*"
        ])
        
        return "\n".join(lines)


class JSONGenerator(BaseFormatGenerator):
    """Generate JSON reports."""
    
    async def generate(self, content: Dict[str, Any], output_filename: str) -> str:
        """Generate JSON report."""
        output_path = self.output_dir / f"{output_filename}.json"
        
        # Convert content to JSON-serializable format
        json_content = await self._prepare_json_content(content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_content, f, indent=2, default=str, ensure_ascii=False)
        
        logger.info(f"JSON report generated: {output_path}")
        return str(output_path)
    
    async def _prepare_json_content(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare content for JSON serialization."""
        json_content = {}
        
        # Metadata
        metadata = content.get('metadata')
        if metadata:
            json_content['metadata'] = {
                'report_id': metadata.report_id,
                'title': metadata.title,
                'report_type': metadata.report_type.value,
                'target_binary': metadata.target_binary,
                'generated_at': metadata.generated_at.isoformat(),
                'generated_by': metadata.generated_by,
                'version': metadata.version,
                'analysis_duration': str(metadata.analysis_duration) if metadata.analysis_duration else None,
                'total_findings': metadata.total_findings,
                'critical_findings': metadata.critical_findings,
                'high_findings': metadata.high_findings,
                'medium_findings': metadata.medium_findings,
                'low_findings': metadata.low_findings,
                'info_findings': metadata.info_findings,
                'ai_models_used': metadata.ai_models_used,
                'data_sources': metadata.data_sources
            }
        
        # Content sections
        json_content['executive_summary'] = content.get('executive_summary', '')
        json_content['risk_assessment'] = content.get('risk_assessment', '')
        json_content['recommendations'] = content.get('recommendations', '')
        
        # Findings
        findings = content.get('findings', [])
        json_content['findings'] = []
        for finding in findings:
            finding_dict = {
                'id': finding.id,
                'title': finding.title,
                'description': finding.description,
                'severity': finding.severity.value,
                'category': finding.category,
                'evidence': finding.evidence,
                'recommendations': finding.recommendations,
                'confidence': finding.confidence,
                'cvss_score': finding.cvss_score,
                'cve_ids': finding.cve_ids,
                'affected_components': finding.affected_components,
                'exploitation_difficulty': finding.exploitation_difficulty,
                'business_impact': finding.business_impact,
                'ai_generated': finding.ai_generated,
                'ai_confidence': finding.ai_confidence
            }
            json_content['findings'].append(finding_dict)
        
        # AI Insights
        ai_insights = content.get('ai_insights', [])
        json_content['ai_insights'] = []
        for insight in ai_insights:
            insight_dict = {
                'insight_type': insight.insight_type,
                'title': insight.title,
                'content': insight.content,
                'confidence': insight.confidence,
                'model_used': insight.model_used,
                'generation_time': insight.generation_time.isoformat(),
                'supporting_evidence': insight.supporting_evidence,
                'related_findings': insight.related_findings
            }
            json_content['ai_insights'].append(insight_dict)
        
        # Analysis context
        json_content['analysis_context'] = content.get('analysis_context', {})
        
        return json_content


class DOCXGenerator(BaseFormatGenerator):
    """Generate DOCX reports using python-docx."""
    
    async def generate(self, content: Dict[str, Any], output_filename: str) -> str:
        """Generate DOCX report."""
        if not DOCX_AVAILABLE:
            # Fallback to Markdown
            markdown_gen = MarkdownGenerator(str(self.output_dir))
            markdown_path = await markdown_gen.generate(content, output_filename)
            logger.warning(f"DOCX generation not available, generated Markdown instead: {markdown_path}")
            return markdown_path
        
        output_path = self.output_dir / f"{output_filename}.docx"
        
        # Create document
        doc = Document()
        
        # Add styles
        self._setup_docx_styles(doc)
        
        # Add content
        await self._add_docx_header(doc, content)
        await self._add_docx_executive_summary(doc, content)
        await self._add_docx_findings_overview(doc, content)
        await self._add_docx_detailed_findings(doc, content)
        await self._add_docx_ai_insights(doc, content)
        await self._add_docx_recommendations(doc, content)
        
        # Save document
        doc.save(str(output_path))
        
        logger.info(f"DOCX report generated: {output_path}")
        return str(output_path)
    
    def _setup_docx_styles(self, doc):
        """Setup custom styles for DOCX document."""
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = docx.shared.Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Heading styles
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = docx.shared.Pt(16)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(52, 73, 94)  # #34495e
    
    async def _add_docx_header(self, doc, content: Dict[str, Any]):
        """Add header to DOCX document."""
        metadata = content.get('metadata')
        
        if metadata:
            title = doc.add_paragraph(metadata.title, style='CustomTitle')
            doc.add_paragraph(f"Generated: {self._format_timestamp(metadata.generated_at)}")
            doc.add_paragraph(f"Target: {Path(metadata.target_binary).name}")
        else:
            title = doc.add_paragraph("Security Analysis Report", style='CustomTitle')
        
        doc.add_page_break()
    
    async def _add_docx_executive_summary(self, doc, content: Dict[str, Any]):
        """Add executive summary to DOCX."""
        doc.add_heading('Executive Summary', level=1)
        summary = content.get('executive_summary', 'No executive summary available.')
        doc.add_paragraph(summary)
    
    async def _add_docx_findings_overview(self, doc, content: Dict[str, Any]):
        """Add findings overview to DOCX."""
        doc.add_heading('Findings Overview', level=1)
        
        findings = content.get('findings', [])
        if not findings:
            doc.add_paragraph("No findings to report.")
            return
        
        # Create severity summary table
        severity_counts = {}
        for severity in ReportSeverity:
            severity_counts[severity] = len([f for f in findings if f.severity == severity])
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'Count'
        
        # Data rows
        for severity, count in severity_counts.items():
            if count > 0:
                row_cells = table.add_row().cells
                row_cells[0].text = severity.value.title()
                row_cells[1].text = str(count)
        
        doc.add_paragraph(f"Total Findings: {len(findings)}")
    
    async def _add_docx_detailed_findings(self, doc, content: Dict[str, Any]):
        """Add detailed findings to DOCX."""
        doc.add_heading('Detailed Findings', level=1)
        
        findings = content.get('findings', [])
        if not findings:
            doc.add_paragraph("No detailed findings to report.")
            return
        
        # Sort findings by severity
        sorted_findings = sorted(findings, key=lambda f: self._get_severity_priority(f.severity))
        
        for finding in sorted_findings:
            doc.add_heading(f"{finding.title} ({finding.severity.value.upper()})", level=2)
            
            doc.add_paragraph(f"Category: {finding.category}")
            doc.add_paragraph(f"Description: {finding.description}")
            
            if finding.evidence:
                doc.add_paragraph("Evidence:")
                for evidence in finding.evidence:
                    p = doc.add_paragraph(evidence, style='List Bullet')
            
            if finding.recommendations:
                doc.add_paragraph("Recommendations:")
                for rec in finding.recommendations:
                    p = doc.add_paragraph(rec, style='List Bullet')
            
            doc.add_paragraph(
                f"Confidence: {finding.confidence:.1%} | "
                f"Exploitation Difficulty: {finding.exploitation_difficulty} | "
                f"Business Impact: {finding.business_impact}"
            )
    
    async def _add_docx_ai_insights(self, doc, content: Dict[str, Any]):
        """Add AI insights to DOCX."""
        doc.add_heading('AI-Generated Insights', level=1)
        
        ai_insights = content.get('ai_insights', [])
        if not ai_insights:
            doc.add_paragraph("No AI insights available.")
            return
        
        for insight in ai_insights:
            doc.add_heading(insight.title, level=2)
            doc.add_paragraph(insight.content)
            doc.add_paragraph(f"Confidence: {insight.confidence:.1%} | Model: {insight.model_used}")
    
    async def _add_docx_recommendations(self, doc, content: Dict[str, Any]):
        """Add recommendations to DOCX."""
        doc.add_heading('Recommendations', level=1)
        recommendations = content.get('recommendations', 'No recommendations available.')
        doc.add_paragraph(recommendations)


class TextGenerator(BaseFormatGenerator):
    """Generate plain text reports."""
    
    async def generate(self, content: Dict[str, Any], output_filename: str) -> str:
        """Generate plain text report."""
        output_path = self.output_dir / f"{output_filename}.txt"
        
        text_content = await self._create_text_content(content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        logger.info(f"Text report generated: {output_path}")
        return str(output_path)
    
    async def _create_text_content(self, content: Dict[str, Any]) -> str:
        """Create complete text content."""
        lines = []
        
        # Header
        lines.extend([
            "=" * 80,
            "INTELLICRACK SECURITY ANALYSIS REPORT",
            "=" * 80
        ])
        
        metadata = content.get('metadata')
        if metadata:
            lines.extend([
                f"Title: {metadata.title}",
                f"Generated: {self._format_timestamp(metadata.generated_at)}",
                f"Target: {Path(metadata.target_binary).name}",
                f"Report ID: {metadata.report_id}",
                ""
            ])
        
        # Executive Summary
        lines.extend([
            "EXECUTIVE SUMMARY",
            "-" * 40,
            content.get('executive_summary', 'No executive summary available.'),
            ""
        ])
        
        # Findings Overview
        lines.extend([
            "FINDINGS OVERVIEW",
            "-" * 40
        ])
        
        findings = content.get('findings', [])
        if findings:
            severity_counts = {}
            for severity in ReportSeverity:
                severity_counts[severity] = len([f for f in findings if f.severity == severity])
            
            for severity, count in severity_counts.items():
                if count > 0:
                    lines.append(f"{severity.value.title()}: {count}")
            
            lines.extend([f"Total Findings: {len(findings)}", ""])
        else:
            lines.extend(["No findings to report.", ""])
        
        # AI Insights
        lines.extend([
            "AI-GENERATED INSIGHTS",
            "-" * 40
        ])
        
        ai_insights = content.get('ai_insights', [])
        if ai_insights:
            for i, insight in enumerate(ai_insights, 1):
                lines.extend([
                    f"{i}. {insight.title}",
                    f"   {insight.content}",
                    f"   Confidence: {insight.confidence:.1%} | Model: {insight.model_used}",
                    ""
                ])
        else:
            lines.extend(["No AI insights available.", ""])
        
        # Detailed Findings
        lines.extend([
            "DETAILED FINDINGS",
            "-" * 40
        ])
        
        if findings:
            sorted_findings = sorted(findings, key=lambda f: self._get_severity_priority(f.severity))
            
            for i, finding in enumerate(sorted_findings, 1):
                lines.extend([
                    f"{i}. {finding.title} ({finding.severity.value.upper()})",
                    f"   Category: {finding.category}",
                    f"   Description: {finding.description}"
                ])
                
                if finding.evidence:
                    lines.append("   Evidence:")
                    for evidence in finding.evidence:
                        lines.append(f"     - {evidence}")
                
                if finding.recommendations:
                    lines.append("   Recommendations:")
                    for rec in finding.recommendations:
                        lines.append(f"     - {rec}")
                
                lines.extend([
                    f"   Confidence: {finding.confidence:.1%} | "
                    f"Exploitation Difficulty: {finding.exploitation_difficulty} | "
                    f"Business Impact: {finding.business_impact}",
                    ""
                ])
        else:
            lines.extend(["No detailed findings to report.", ""])
        
        # Risk Assessment
        lines.extend([
            "RISK ASSESSMENT",
            "-" * 40,
            content.get('risk_assessment', 'No risk assessment available.'),
            ""
        ])
        
        # Recommendations
        lines.extend([
            "RECOMMENDATIONS",
            "-" * 40,
            content.get('recommendations', 'No recommendations available.'),
            ""
        ])
        
        # Footer
        lines.extend([
            "=" * 80,
            f"Report generated by Intellicrack AI Report Generator",
            f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}",
            "=" * 80
        ])
        
        return "\n".join(lines)


# Format generator factory
class FormatGeneratorFactory:
    """Factory for creating format generators."""
    
    @staticmethod
    def create_generator(format_type: ReportFormat, output_dir: str = "reports") -> BaseFormatGenerator:
        """Create a format generator for the specified type."""
        generators = {
            ReportFormat.HTML: HTMLGenerator,
            ReportFormat.PDF: PDFGenerator,
            ReportFormat.MARKDOWN: MarkdownGenerator,
            ReportFormat.JSON: JSONGenerator,
            ReportFormat.DOCX: DOCXGenerator,
            ReportFormat.TXT: TextGenerator
        }
        
        generator_class = generators.get(format_type)
        if not generator_class:
            raise ValueError(f"Unsupported format: {format_type}")
        
        return generator_class(output_dir)
    
    @staticmethod
    def get_supported_formats() -> List[ReportFormat]:
        """Get list of supported report formats."""
        return list(ReportFormat)


# Export the main components
__all__ = [
    'BaseFormatGenerator',
    'HTMLGenerator',
    'PDFGenerator', 
    'MarkdownGenerator',
    'JSONGenerator',
    'DOCXGenerator',
    'TextGenerator',
    'FormatGeneratorFactory'
]